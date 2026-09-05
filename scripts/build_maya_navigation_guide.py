from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Maya_Navigation_Guide.docx"

INK = "111815"
MUTED = "5E6862"
LINE = "D9D9D9"
PALE = "F4F5F1"
PALE_GREEN = "F3F7E7"
LIME = "D7FF38"
WHITE = "FFFFFF"
ORANGE = "FF6B35"


def set_font(run, name: str = "Arial", size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = LINE, size: str = "6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_header_and_bands(table):
    repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_fill(cell, INK)
            elif row_index % 2 == 0:
                set_cell_fill(cell, PALE)


def set_widths(table, widths: list[float]):
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def style_table_text(table, body_size: float = 8.3):
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                if cell_index == 0 and row_index > 0:
                    paragraph.paragraph_format.keep_with_next = False
                for run in paragraph.runs:
                    set_font(run, size=8 if row_index == 0 else body_size, bold=True if row_index == 0 else None, color=WHITE if row_index == 0 else INK)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, body_size: float = 8.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    set_table_borders(table)
    set_repeat_header_and_bands(table)
    if widths:
        set_widths(table, widths)
    style_table_text(table, body_size=body_size)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)
    return table


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2F5B50")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    new_run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])
    set_font(run, size=8, color=MUTED)


def add_heading(doc, title: str, level: int = 1, intro: str | None = None):
    paragraph = doc.add_heading(title, level=level)
    paragraph.paragraph_format.keep_with_next = True
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(10)
    return paragraph


def remove_style_border(style):
    paragraph_properties = style._element.get_or_add_pPr()
    paragraph_border = paragraph_properties.find(qn("w:pBdr"))
    if paragraph_border is not None:
        paragraph_properties.remove(paragraph_border)


def remove_paragraph_border(paragraph):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    paragraph_border = paragraph_properties.find(qn("w:pBdr"))
    if paragraph_border is not None:
        paragraph_properties.remove(paragraph_border)


def add_label(doc, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text.upper())
    set_font(run, name="Arial", size=8.5, bold=True, color=ORANGE)
    return paragraph


def add_bullets(doc, items: list[str], numbered: bool = False):
    if numbered:
        for index, item in enumerate(items, start=1):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.24)
            paragraph.paragraph_format.first_line_indent = Inches(-0.24)
            paragraph.paragraph_format.space_after = Pt(4)
            set_font(paragraph.add_run(f"{index}.  "), bold=True)
            paragraph.add_run(item)
        return
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_code_block(doc, lines: list[str]):
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(2)
        set_font(paragraph.add_run(line), name="Consolas", size=8.2, color=INK)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title.font.size = Pt(31)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("000000")
    title.paragraph_format.space_after = Pt(12)
    remove_style_border(title)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor.from_string("000000")
    subtitle.font.italic = False

    for name, size in (("Heading 1", 21), ("Heading 2", 15), ("Heading 3", 11.5)):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("000000")
        style.paragraph_format.space_before = Pt(12 if name != "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(9.7)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.14


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(3)
    set_font(paragraph.add_run("Maya Workspace  |  Synthetic data  |  Page "), size=8, color=MUTED)
    add_page_number(paragraph)


def page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    add_footer(section)

    core_properties = doc.core_properties
    core_properties.title = "Maya Workspace Navigation Guide"
    core_properties.subject = "How to navigate Maya Workspace and its project-scoped capabilities"
    core_properties.author = "Kearney"
    core_properties.keywords = "Maya, supply chain, navigation, projects, operations research, evidence, agents"

    add_label(doc, "Workspace user guide")
    title = doc.add_paragraph(style="Title")
    title.add_run("Maya Workspace Navigation Guide")
    remove_paragraph_border(title)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Projects  Mounted work  Data and graph  Evidence")
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(18)
    meta.paragraph_format.space_after = Pt(16)
    set_font(meta.add_run("Version 5  |  5 September 2026  |  Front-end concept"), name="Consolas", size=9, bold=True, color=MUTED)
    opening = doc.add_paragraph()
    opening.paragraph_format.space_after = Pt(13)
    set_font(opening.add_run("This guide explains the complete Maya Workspace journey: choose a client and project, bring in governed data and IoT signals, navigate mounted specialist apps, work in Playground, inspect accountable agents and team members, compare an operations-research response, and stop at a human decision gate."), size=12)

    add_heading(doc, "Use this guide", level=2)
    add_bullets(doc, [
        "For a five minute walkthrough, follow the Fastest demonstration route on the next page.",
        "For project delivery, start with Workspace hierarchy and Project surfaces.",
        "For trust and governance, read Evidence receipts, Playground and accountability, and Concept boundaries before demonstrating an approval or optimization flow.",
    ], numbered=True)
    p = doc.add_paragraph()
    run = p.add_run("Important boundary  ")
    set_font(run, bold=True)
    p.add_run("All companies, records, calculations, alerts, maps, agents, source connections, and outcomes in the current experience are synthetic. Browser interactions create session-only state. No live enterprise system, external data provider, solver, credential, model training loop, or operational write-back is connected.")

    page_break(doc)
    add_label(doc, "Orientation")
    add_heading(doc, "Platform map", intro="Maya organizes work around a client project. The platform, apps, agents, and evidence all inherit that project context.")
    add_code_block(doc, [
        "Workspace",
        "  Clients and projects",
        "    By client: Client > Tower > Project",
        "    By tower: Tower > Client > Project",
        "      Project context bar",
        "        Apps and mounted app shortcuts",
        "        Agent and team accountability launchers",
        "      Project workspace",
        "        Overview and decisions",
        "        Data and graph",
        "        Playground sessions",
        "        Evidence and controls",
        "  Operations World (Global or Regional)",
    ])
    add_heading(doc, "Fastest demonstration route", level=2)
    add_table(doc, ["Step", "Click", "What the audience should notice", "Terminal result"], [
        ["1", "Workspace", "The outer left rail contains the only project hierarchy. It opens By client and can switch to By tower without changing the selected project.", "Client and project workspace"],
        ["2", "Any project KPI", "Every headline value exposes source type, formula, inputs, variable ID, confidence, trace, reviewer, and access scope.", "Evidence drawer"],
        ["3", "Apps in the mounted-work bar", "The strip shows only apps mounted to this project; Apps opens the complete project catalog without adding another project tab.", "Project app catalog"],
        ["4", "MineralAtlas", "Country, reserve, refinery, route, and product context use an atlas workflow rather than a generic dashboard.", "Mineral sourcing scenario"],
        ["5", "Playground", "Open a stable session ID, inspect its message IDs and app calls, steer an active child, and stop at the human gate.", "Session result and trace"],
        ["6", "Data & graph", "Uploads, IoT requests, data search, source records, and knowledge graph inspection stay in one project surface.", "Dataset, connector, or evidence receipt"],
        ["7", "Decisions", "High-level choices decompose into lower-level decisions, variables, methods, and evidence.", "Expert review draft"],
        ["8", "Operations World", "Toggle Global or Regional, choose a region, inspect the map, and start a dependency, route, or value intake from retained context.", "Project intake draft"],
    ], [0.42, 1.35, 3.25, 1.45], body_size=8.0)

    page_break(doc)
    add_label(doc, "Client delivery")
    add_heading(doc, "Workspace hierarchy", intro="Workspace is the application root. The outer project rail can organize the same records by client or by tower. Every selected project owns its data, apps, agents, decisions, collaborators, evidence, and controls.")
    add_table(doc, ["Level", "Purpose", "Isolation rule", "Primary actions"], [
        ["Sector", "Groups related clients and reusable domain patterns.", "Client records remain separated even when sector methods are shared.", "Filter portfolio, compare common risks, reuse approved templates"],
        ["Client", "Represents the governed tenant and commercial relationship.", "Client data, identities, residency, and policies must be enforced server-side in production.", "Open client projects, manage access, inspect client-wide evidence"],
        ["Project", "Contains the decision memory for one engagement or operating problem.", "Uploads, mounted apps, traces, decisions, and receipts carry the project key.", "Work the problem, collaborate, review evidence, prepare decisions"],
        ["Decision", "Packages a specific choice and its lifecycle.", "The decision retains variables, methods, scenarios, owners, approvals, evidence, and outcome measures.", "Validate, simulate, approve, execute, measure"],
        ["Evidence receipt", "Explains a displayed claim or action.", "A receipt fails closed when its reference is unknown or belongs to another project.", "Inspect source, formula, inputs, version, confidence, access"],
    ], [0.72, 1.68, 2.65, 1.5], body_size=8.2)
    add_heading(doc, "How to switch projects", level=2)
    add_bullets(doc, [
        "Use the Project path rail. By client expands Client > Tower > Project; By tower expands Tower > Client > Project. The toggle changes the grouping, not the active project.",
        "The hierarchy is rendered once, in the outer rail. Project pages do not repeat a second client/tower/project tree inside the working canvas.",
        "Search inside the project rail by client, tower, project name, project code, or problem. Matching paths expand automatically.",
        "New client, New project, and Collapse sit together at the top of the rail. Collapse with that control or Command B. On a narrow screen, use the menu button to open the rail as a drawer.",
        "Read the top breadcrumb for the complete Client, Tower, Project path. The bar below identifies the current surface, mounted apps, accountable agents and team members, and the active Playground session.",
        "Use Search workspace or Command K to jump directly to a project or to a capability inside the selected project.",
        "The URL records the project ancestry, surface, mounted app studio, Playground session, and app run when those records are visible. Browser Back and Forward restore the same valid project-scoped route.",
        "Opening a specialist application keeps one return point. Use Back to project in the context bar or Browser Back to return to the project tab that launched it.",
        "Each project preserves its own browser-session app mounts, uploads, connector requests, agent messages, traces, reruns, draft agents, and team assignments while the page remains open.",
    ])
    page_break(doc)
    add_heading(doc, "Onboard a client and project", level=2)
    add_bullets(doc, [
        "Choose New client to register a session-only client under a sector, with a named client lead and engagement context.",
        "Choose New project to select the client, define the problem and outcome, and assign one client collaborator plus one Kearney specialist.",
        "A new project starts empty: zero datasets, apps, agents, decisions, graph nodes, and runs until the team adds them explicitly.",
        "Creation ends at a project overview and a browser-session receipt; no tenant, identity, cloud resource, or production record is provisioned.",
    ], numbered=True)

    add_label(doc, "Synthetic portfolio")
    add_heading(doc, "Ten client projects", intro="The portfolio is designed to demonstrate different sectors, decision problems, data shapes, expert roles, and app combinations.")
    add_table(doc, ["Code", "Sector and client", "Project", "Problem demonstrated"], [
        ["P-001", "Mobility and EV\nApex Mobility", "Anode Shield", "Graphite concentration, port delay, qualification, launch service, and margin exposure"],
        ["P-002", "Life Sciences\nHelixora Therapeutics", "Cold Chain Promise", "Batch release, refrigerated storage, scarce packaging, patient priority, and stability windows"],
        ["P-003", "Semiconductors\nOrionSilicon Foundry", "Fab Recovery X9", "Tool failure, utilities, chemicals, work in process, qualification, and customer allocation"],
        ["P-004", "Food and Agriculture\nVerdant Foods Cooperative", "Harvest to Shelf", "Seasonality, perishability, farm supply, cold slots, labor, and food waste"],
        ["P-005", "Aerospace\nStratos Aero Systems", "Forging Continuity", "Certified titanium forgings, opaque subtiers, long lead times, and second-source qualification"],
        ["P-006", "Energy and Grid\nSolara Gridworks", "Copper and Rare Earth Portfolio", "Mineral price, refining concentration, capacity, carbon, cash, and portfolio choice"],
        ["P-007", "Mining and Critical Minerals\nTerraMetals Alliance", "Lithium to Cell Provenance", "Mine, refinery, offtake, water, permit, rights, origin, and traceability constraints"],
        ["P-008", "Ports and Maritime\nBlueHarbor Ports and Cargo", "Global Berth to Door", "Berths, vessels, containers, customs, labor, secure handoffs, and inland transfers"],
        ["P-009", "Industrial Automation\nTitanWorks Robotics", "Factory and Service Continuity", "Machine reliability, scarce skills, production, service spares, and installed-base uptime"],
        ["P-010", "Retail and Ecommerce\nMeridian Commerce Group", "Omnichannel Peak", "Promotion demand, available to promise, picking labor, fleet capacity, last mile, and margin"],
    ], [0.48, 1.62, 1.5, 3.05], body_size=7.9)
    p = doc.add_paragraph()
    set_font(p.add_run("Demonstration note  "), bold=True)
    p.add_run("Each project uses synthetic scale, metric, app, agent, method, and variable records. These records show how a production project could be structured; they do not represent the named organizations.")

    page_break(doc)
    add_label(doc, "Project workspace")
    add_heading(doc, "Project surfaces", intro="Five compact tabs stay in the same position and size as the user changes tasks. The tab grid adapts to the available width without a horizontal scroller; apps and people are launched from the mounted-work bar above it.")
    add_table(doc, ["Tab", "Use it for", "Important interactions", "End state"], [
        ["Overview", "Read outcome, project KPIs, mounted-work summary, decisions, and recent session history.", "Open KPI evidence, a decision, an app, a person, a prior session, or project data.", "Evidence drawer or exact work record"],
        ["Decisions", "Decompose the project outcome into D0 through D3 decisions.", "Select a decision, open its evidence, inspect L2 L1 L0 variables, open method references, create review draft.", "Expert review draft receipt"],
        ["Data & graph", "Search, stage, and inspect project-owned sources and their connected entities in one place.", "Upload or request data, query files and records, switch to Knowledge graph, select a node, and open evidence.", "Dataset, connector, graph, or evidence receipt"],
        ["Playground", "Continue agent-first work in a full session workspace.", "Open or create a session, enter a prompt, advance or cancel trace, steer, inspect app calls, and reopen immutable results.", "Human gate, session result, or draft agent manifest"],
        ["Controls", "Review concept controls and production gaps.", "Inspect evidence, access, control status, and boundary statements.", "Control receipt"],
    ], [1.0, 2.2, 2.65, 1.0], body_size=7.8)
    add_heading(doc, "Mounted work bar", level=2)
    add_table(doc, ["Launcher", "Purpose", "Where it opens"], [
        ["Apps", "Open the full app dependency and mount catalog for this project.", "The project workspace"],
        ["Mounted app chips", "Jump directly to the apps this project actually uses. The active app remains visible; the overflow count opens Apps.", "A distinct app studio with the same project frame"],
        ["Agents", "See project agents and who is running which work. Compact names expose the most relevant identities.", "One right-side accountability panel"],
        ["Team", "See client and Kearney collaborators, their project rights, and attributed activity.", "The same right-side accountability panel"],
        ["Playground session", "Return to the current agent conversation or session history.", "Playground in the project workspace"],
    ], [1.35, 3.8, 1.75], body_size=8.1)
    p = doc.add_paragraph("The Agents and Team launchers are mutually exclusive: opening one identity replaces the other in the right panel. Each profile shows only explicitly attributed project activity. Production attribution requires authenticated identities and signed, durable audit events.")
    add_heading(doc, "What project persistence means", level=2)
    p = doc.add_paragraph("The concept preserves per-project state in browser memory only. The signed-in Maya Rao fixture is one Kearney portfolio collaborator with an explicit membership in each seeded project. Project entry and governed session mutations evaluate that identity's declared capability grants; the interface never borrows a client owner's rights. A production implementation still needs authenticated tenant context, server-side project authorization, encrypted object storage, a project registry, lineage storage, versioned event history, retention policies, and audit controls. Do not interpret session restoration as backend persistence.")

    add_label(doc, "Interaction completeness")
    add_heading(doc, "Clickable surfaces and terminal states", intro="Use this section to explain what a control does and how the user knows that it completed.")
    add_table(doc, ["Control family", "Expected behavior", "Visible completion", "No silent claim"], [
        ["Navigation and search", "Open Workspace, Operations World, a project, tab, app, decision, graph, or agent surface and update the URL.", "Destination heading receives focus and Back restores the route.", "No placeholder page"],
        ["Filters and selectors", "Recalculate or filter the deterministic fixture shown on screen.", "Selection, counts, chart, map, list, or inspector changes.", "No external refresh"],
        ["Metric and evidence controls", "Open a project evidence drawer or session action receipt.", "Receipt names the value, basis, context, and boundary.", "No unsupported source access"],
        ["Create and update controls", "Change browser-session state only when the user has enough context.", "Saved or Completed receipt, updated local state, and receipt ledger entry.", "No backend persistence"],
        ["Blocked controls", "Stay disabled or return a specific reason when prerequisites are missing.", "Blocked receipt or visible gate explanation.", "No pretend success"],
        ["Future connectors", "Open a preview or contract rather than simulate a hidden integration.", "Future adapter explanation and missing production requirements.", "No credential or service created"],
        ["Approval and release", "Apply the decision lifecycle gate and named human authority.", "Session-only stage change or an exact blocker.", "No ERP TMS WMS or payment write-back"],
    ], [1.15, 2.35, 2.15, 1.25], body_size=8.0)
    add_heading(doc, "Session receipt ledger", level=2)
    add_bullets(doc, [
        "Any outcome-producing control creates a receipt with status, artifact ID, recorded context, timestamp, and execution boundary.",
        "Open the user menu and choose Open session receipt ledger to revisit the most recent receipt.",
        "The ledger retains up to 24 entries in the current browser session. It is not a production audit store.",
    ])

    page_break(doc)
    add_label(doc, "Project decisions")
    add_heading(doc, "From signal to governed decision", intro="Decisions and review live inside a project. They connect app evidence to a governed choice, human authority, execution package, and measurable outcome.")
    add_table(doc, ["Stage", "Question", "Required content", "Exit condition in the concept"], [
        ["Detect", "What changed and why might it matter", "Signal, affected entity, horizon, confidence, initial value exposure", "Named decision exists"],
        ["Validate", "Is the signal credible and relevant", "Evidence records, source references, data quality, graph path, project variable contract", "Evidence threshold reached"],
        ["Simulate", "What responses are feasible under uncertainty", "Scenarios, constraints, assumptions, method stack, service, cost, cash, risk, carbon", "Candidate inside hard envelope"],
        ["Approve", "Who can accept the tradeoff", "Recommendation, evidence, finance review, decision rights, fallback", "Named human approval"],
        ["Execute", "What work packages should be released", "Tasks, owners, timing, interfaces, status, rollback or fallback", "Session execution state"],
        ["Measure", "Did the decision produce the expected outcome", "Baseline, target, realized value, measurement window, overrides, failure modes", "Outcome receipt"],
    ], [0.72, 1.6, 3.45, 1.08], body_size=8.0)
    add_heading(doc, "Decision tree and variable taxonomy", level=2)
    p = doc.add_paragraph("The D0 to D3 decision hierarchy is different from the supply chain variable taxonomy. D0 to D3 expresses choice decomposition. L2, L1, and L0 identify drivers and measurable variables. A decision node references the relevant taxonomy pack and OR methods, then points to evidence. The current project packs use canonical identifiers from the taxonomy rather than invented labels.")
    add_code_block(doc, [
        "D0 project outcome decision",
        "  D1 portfolio or policy decision",
        "    D2 sourcing production logistics or demand decision",
        "      D3 executable choice",
        "        L2 driver family",
        "          L1 subdriver",
        "            L0 measurable variable",
        "              calculation and evidence receipt",
    ])

    page_break(doc)
    add_label(doc, "Specialist applications")
    add_heading(doc, "Ten distinct app workspaces", intro="Each application has its own visual grammar and operating rhythm. All apps use the same project graph and evidence contracts, but they do not look or behave like interchangeable dashboards.")
    add_table(doc, ["Application", "Visual and interaction model", "Primary job", "Output"], [
        ["RiskRadar", "Coral signal room with propagation bubbles, criticality filters, and causal change ledger", "Trace a disruption through suppliers, materials, plants, products, orders, and margin", "Risk control brief"],
        ["Network Optimizer", "Lime formulation editor, method library, scenario lab, Pareto view, run console, and release gate", "Frame and compare a feasible sourcing, capacity, inventory, or logistics response", "Versioned solution package"],
        ["FlowLens", "Cyan material-to-cash waterfall and action queue", "Show where inventory, transit, terms, or delays trap cash and margin", "Cash action package"],
        ["DemandSense", "Violet forecast fan, causal notebook, scenario controls, and gap table", "Commit an explainable demand range and expose supply gaps", "Demand contract"],
        ["SupplierGraph", "Green n-tier dependency graph and qualification funnel", "Find dependency, ownership, capability, alternative source, and evidence gaps", "Qualified shortlist"],
        ["MineralAtlas", "Copper reserve-to-refinery world atlas with mineral and country selectors", "Trace critical mineral concentration from country to product", "Mineral sourcing scenario"],
        ["WorkforceStudio", "Indigo skills matrix and shift builder", "Match qualified people to constrained work and expose certification gaps", "Workforce capacity plan"],
        ["ManufacturingTwin", "Steel-blue plant flow, constraint timeline, and schedule controls", "Schedule bottlenecks, maintenance, yield, and production commitments", "Production plan"],
        ["LogisticsRadar", "Teal radar map and transfer playback", "Control cargo, route, customs, custody, and last-mile handoffs", "Route and handoff plan"],
        ["QualityGenealogy", "Amber batch genealogy and release gate", "Trace lot, process, certificate, deviation, release, and customer consequence", "Release or containment plan"],
    ], [1.18, 2.8, 2.3, 1.1], body_size=7.8)

    page_break(doc)
    add_label(doc, "Using a specialist app")
    add_heading(doc, "How to inspect an app", intro="Each app opens on the specialist's task, with concise controls and inspectable methods, data, measures, and handoffs rather than a presentation layer.")
    add_table(doc, ["Section", "What to inspect", "Trace action"], [
        ["Task surface", "Decision scope, horizon, workflow state, accountable roles, and current constraint", "Use the app-specific controls and decision links"],
        ["Methods", "Purpose, model family, formulation, and validation contract for each app method", "Open method reference receipt"],
        ["Data and controls", "Data grain, expected source, target freshness, fitness checks, policy rule, owner, and evidence", "Open data contract or control receipt"],
        ["Measure and hand off", "KPI definition, target, accountable owner, downstream destination, trigger, artifact, and limitation", "Open KPI or handoff receipt"],
        ["Change and learning", "Before, current, forecast, delta, explanation, confidence, trigger, owner, and expected outcome feedback", "Trace selected change"],
    ], [1.35, 4.6, 1.35], body_size=8.2)
    add_heading(doc, "Mounted app dependency graph", level=2)
    p = doc.add_paragraph("The app graph explains how applications cooperate without collapsing into one interface. DemandSense can inform Network Optimizer. SupplierGraph and MineralAtlas can challenge sourcing feasibility. ManufacturingTwin and WorkforceStudio can constrain capacity. LogisticsRadar can validate route execution. QualityGenealogy can block release. RiskRadar can propagate external and operational change. FlowLens can translate a candidate into cash and margin. Every edge should identify the project, source app, target app, shared entity keys, variable IDs, evidence version, and handoff status in production.")
    add_bullets(doc, [
        "Open a mounted app chip for a direct jump, or choose Apps in the mounted-work bar to inspect and change the project catalog.",
        "Use Mount or Unmount to change the session manifest. Core project views remain available even when an app is unmounted.",
        "Open the mount manifest receipt to see the fixture count and project boundary.",
        "Use Run history to reopen an exact immutable run ID, inspect its report, trace, parent, inputs, and outputs, or open the application home for a new task.",
        "Editable numeric assumptions remain visible when invalid. Correct the inline range or step error before Rerun with changes becomes available.",
    ])

    page_break(doc)
    add_label(doc, "Playground")
    add_heading(doc, "Work with specialist agents", intro="Playground is the agent-first work surface. Its code-style layout exposes the conversation, stable message IDs, visible product trace, app calls, evidence, results, steering, and human decision gates without changing the surrounding project frame.")
    add_table(doc, ["Area", "What it shows", "How to use it", "Boundary"], [
        ["Session history", "Stable session ID, title, status, agent, activity, and result", "Continue a previous task or start a new project-scoped conversation", "History is browser-session fixture data"],
        ["Run-as control", "The specialist assigned to the next Playground prompt", "Choose an existing project agent or open New agent to draft a manifest", "Profiles are synthetic"],
        ["Conversation", "User prompt, agent response, and system status", "Ask a project-specific question. The deterministic response binds to the current project and selected specialist", "No model API is called"],
        ["Tool trace", "Step, agent, activity, detail, and graph node", "Start, advance, replay, cancel, and inspect the visible execution trace", "This is not hidden chain of thought"],
        ["Visual steering", "Instructions such as pin evidence, require human review, preserve service floor, or exclude an entity", "Click a steering control and confirm the instruction appears in the trace context", "No policy or model weights change"],
        ["Experience profile", "Level, years, evaluated runs, approved runs, calibration, override rate, failure rate, skills, MCPs, tools, and authority", "Open the profile fixture receipt before relying on a capability", "Statistics are demo fixtures"],
        ["Agent builder", "Name, specialty, selected Skills file name, requested MCPs, and draft tools", "Create and save a draft manifest", "No Skills file is parsed, MCP connected, tool created, evaluation run, or agent deployed"],
    ], [1.05, 2.15, 2.55, 1.2], body_size=7.8)
    add_heading(doc, "Safe agent run sequence", level=2)
    add_bullets(doc, [
        "Open Playground and select the specialist whose authority matches the task, or resume the relevant session from history.",
        "Enter the decision question and press Enter or Send. Use Shift Enter for a new line. Submitting creates or continues an immutable session and starts its visible trace.",
        "Advance the trace and inspect each evidence, graph, variable, formulation, and calculation step. Application-created sessions show Ready until a brief starts the trace.",
        "Add a steering instruction when the project needs an extra constraint or review requirement.",
        "Open evidence receipts for the values that matter to the recommendation.",
        "Complete only at the human gate. Use Cancel when the context is wrong or incomplete.",
    ], numbered=True)
    p = doc.add_paragraph("Agent profiles do not repeat inside Playground. Choose Agents in the mounted-work bar to inspect one accountability record in the right panel. Exactly named actor events are attributed to that identity; sessions and app runs are labeled as involvement unless the fixture has a direct actor record.")

    page_break(doc)
    add_label(doc, "Evidence")
    add_heading(doc, "Trace every displayed claim", intro="A number is trustworthy only when the reader can see what it means, where it came from, how it was produced, and whether it is safe to use for the current decision.")
    add_table(doc, ["Receipt field", "Question answered", "Current concept behavior"], [
        ["Claim and displayed value", "What exactly is being asserted", "Matches the clicked KPI, studio metric, graph fact, decision value, or action result"],
        ["State", "Was it observed, corroborated, inferred, simulated, or proposed", "Synthetic project receipts use Simulated unless a more precise fixture state applies"],
        ["Source kind and locator", "Which upload, SQL table, API, web page, calculation, or fixture supports it", "Uses a bounded fixture locator and says when no source is connected"],
        ["As of and valid for", "When was it true and for which period or decision window", "Uses a fixed concept timestamp and validity statement"],
        ["Version and fingerprint", "Which data or model version produced it", "Uses a fixture version and evidence fingerprint, not a cryptographic content hash claim"],
        ["Formula and inputs", "How was the value derived", "Lists the illustrative calculation and referenced input IDs"],
        ["Variable ID and grain", "Which taxonomy variable and business grain does it represent", "Uses project-specific L2 L1 L0 packs and a stated synthetic grain"],
        ["Confidence and quality", "How certain is it and what quality warnings apply", "Fails closed for unknown or foreign-project references"],
        ["Trace agent and reviewer", "Which worker produced it and who reviewed it", "Names a synthetic playback agent and shows Unreviewed when no review exists"],
        ["Access", "Who may see the evidence", "Shows the client and project boundary; backend enforcement remains future work"],
    ], [1.45, 2.25, 3.25], body_size=7.9)
    add_heading(doc, "Where receipts appear", level=2)
    add_bullets(doc, [
        "Overview KPIs and app studio metrics open the full evidence drawer.",
        "Decision evidence, scenarios, outcomes, and app method or control facts create a browser-session action receipt.",
        "Data & graph search results, knowledge graph nodes, agent trace steps, and app-run inputs point back to project receipts.",
        "Operations World dependency, route, and value intakes retain the selected entity, frame, scenario, and evidence key when starting a project.",
        "Unknown evidence IDs show Not found with zero confidence. A receipt from a different project is not reused.",
    ])

    page_break(doc)
    add_label(doc, "Data and knowledge")
    add_heading(doc, "Use Data & graph", intro="Sources and knowledge relationships share one project surface. A compact mode switch moves between source operations and graph inspection while one query searches the visible project catalog. The concept demonstrates the journey without reading or persisting a selected file.")
    add_table(doc, ["Area", "Use it for", "Terminal result"], [
        ["Project query", "Search Excel, PDF, CSV, JSON, SQL-table, dataset, evidence, connector, variable, and graph fixtures.", "A bounded result opens its source contract, exact evidence receipt, or graph node"],
        ["Sources", "Inspect datasets and connector contracts, stage a local file name, or request IoT and enterprise feeds.", "Dataset or connector receipt"],
        ["Knowledge graph", "Inspect connected project entities and trace context across evidence, variables, calculations, apps, and decisions.", "Selected node, evidence receipt, or Playground steering context"],
    ], [1.35, 4.1, 1.4], body_size=8.1)
    add_heading(doc, "Stage project-owned data", level=2)
    add_table(doc, ["Stage", "User action", "Displayed result", "Production requirement"], [
        ["Select", "Choose a local file or use the provided sample", "Only the file name is captured in browser memory", "Encrypted upload, malware scanning, type limits, tenant authorization"],
        ["Staged", "Confirm the chosen data product", "Project, classification, owner, and proposed dataset name", "Object storage, immutable receipt, consent and retention policy"],
        ["Schema preview", "Advance the demonstration", "Illustrative columns, types, grain, and sample scale", "Actual parser, schema registry, PII detection, quarantine"],
        ["Mapping draft", "Review proposed taxonomy mappings", "Synthetic L0 mappings and expected entity keys", "Human-reviewed semantic mapping and version control"],
        ["Review demo", "Inspect validation and project boundary", "Fixture quality checks and approval gate", "Data steward workflow, access policy, quality service"],
        ["Session receipt", "Complete the demonstration", "Browser-session dataset row and action receipt", "Persistent lineage, catalog registration, audit event"],
    ], [0.88, 1.7, 2.65, 1.72], body_size=8.0)
    page_break(doc)
    add_heading(doc, "Request an IoT or enterprise source", level=2)
    p = doc.add_paragraph("Sources mode includes explicit request templates for operational updates from ports, cargo handlers, warehouses, factories, suppliers, vehicles, robots, and field teams. QR scans, signed custody events, sensor telemetry, robot events, carrier milestones, CDC, and manual exceptions remain request drafts. Policy-review status and fixed-sample replay status are tracked independently, so completing one does not erase the other. Production ingestion must validate identity, device trust, timestamp, location, signature, unit, schema, consent, and access before updating the operational graph.")
    add_table(doc, ["Contributor", "Example event", "Secure contract", "Project use"], [
        ["Port operator", "Berth, gate, customs, seal, or dwell update", "Operator identity, port role, signed timestamp, shipment key", "LogisticsRadar and decision evidence"],
        ["Cargo handler", "QR custody scan or damage exception", "GS1 identifier, device trust, chain-of-custody signature", "QualityGenealogy and LogisticsRadar"],
        ["Supplier", "Capacity, certificate, lead time, or shipment promise", "Supplier tenant, approved portal, evidence attachment, version", "SupplierGraph and RiskRadar"],
        ["Factory or robot", "Production count, state, downtime, quality or consumption event", "Machine identity, event schema, clock quality, plant boundary", "ManufacturingTwin and FlowLens"],
        ["Workforce lead", "Shift, skill, certification, absence, or overtime update", "Role-based access, minimal personal data, regional policy", "WorkforceStudio and Network Optimizer"],
    ], [1.0, 2.0, 2.45, 1.4], body_size=7.9)

    page_break(doc)
    add_label(doc, "Network operations")
    add_heading(doc, "Use Operations World", intro="Operations World is a synthetic operational radar with Global and Regional modes. Its pannable, zoomable map combines geography, routes, assets, cargo, transfers, time, scenarios, and project intake. In Regional mode, the region selector changes the map, KPIs, signals, cash, suppliers, movements, and evidence scope together.")
    add_table(doc, ["Control", "What changes", "What to inspect"], [
        ["Global or Regional", "Switches between the full network and a selected regional operating picture", "Global totals or the selected APAC, Europe, Americas, or Middle East and Africa slice"],
        ["Pan and zoom", "Changes the viewport while preserving filters", "Dense corridors and local handoffs at higher zoom"],
        ["Time frame", "Moves from historical fixture frames through Now and future scenario frames", "Position, arrival state, ETA, delay, cost, volume, and exposure changes"],
        ["Scenario", "Changes trajectory assumptions such as baseline or disruption", "Route geometry, predicted delay, committed quantity, value exposure, and service risk"],
        ["Layers", "Shows or hides ocean, air, road, rail, transfer, asset, cargo, and location records", "Use the smallest layer set that answers the question"],
        ["Hover and keyboard focus", "Shows compact identity, quantity, value, price, status, owner, origin, destination, or personnel details", "No click is required for a quick check"],
        ["Pin selection", "Opens the detailed inspector and retains the canonical entity key", "Evidence, current state, related facts, and project intake context"],
        ["Dependency intake", "Carries the selected dependency, entity ID, frame, scenario, and evidence key into New project", "Creates an empty governed project shell; no case or dataset is fabricated"],
        ["Route intake", "Carries the selected corridor or movement and route-planning context into New project", "Creates a traceable route-intake record"],
        ["Value intake", "Carries the selected entity and value-protection context into New project", "Creates a traceable value-intake record"],
    ], [1.35, 2.45, 3.05], body_size=8.0)
    add_heading(doc, "Map truth boundary", level=2)
    p = doc.add_paragraph("The map uses a local Natural Earth basemap and deterministic route, asset, cargo, transfer, and location fixtures. It does not call Google 3D Tiles, Cesium ion, OpenSky, ADS-B Exchange, AIS, CCTV, traffic, satellite, seismic, or weather services. Production use requires provider terms, attribution, licensing, privacy, retention, and security review in addition to connector engineering.")

    page_break(doc)
    add_label(doc, "Operations research")
    add_heading(doc, "Use the optimization workbench", intro="Network Optimizer separates problem framing from calculation, comparison, validation, and release. The current experience runs a deterministic response calculator; it does not execute a mathematical solver or claim optimality.")
    add_table(doc, ["Workbench area", "Purpose", "Evidence to retain"], [
        ["Decision contract", "Define decision, horizon, selected entity, decision pattern, objective hierarchy, and assumptions", "Project, decision, user, timestamp, variable and constraint version"],
        ["Variable registry", "Bind canonical L0 variables to units, grain, domain, source contract, and decision role", "Variable IDs, bounds, type, unit, missingness, transformations"],
        ["Constraint library", "Classify hard, binding, advisory, and violated constraints", "Expression, tolerance, origin, owner, relaxation authority"],
        ["Method library", "Choose a primary and fallback stack from the 30 handbook methods", "Selection rationale, applicability, validation and fallback"],
        ["Scenario lab", "Represent joint uncertainty, probability, path, severity, and recourse", "Scenario set ID, probabilities, input versions, policy changes"],
        ["Candidate comparison", "Compare cost, service, cash, risk, carbon, and constraint state", "Candidate ID, feasibility, dominance, assumptions, rejected reasons"],
        ["Reproducibility manifest", "Bind the output to the exact context used", "Run ID, model and data version, input fingerprint, method stack, scenario set"],
        ["Release gate", "Require evidence and named human authority before operational handoff", "Approval, scope, expiry, fallback, write-back receipt"],
    ], [1.45, 3.2, 2.2], body_size=8.0)
    add_heading(doc, "Method families in the handbook", level=2)
    add_table(doc, ["Codes", "Family", "Representative use"], [
        ["M-01 to M-04", "Forecasting and uncertainty", "Time series, causal drivers, ensembles, Bayesian updates"],
        ["M-05 to M-08", "Mathematical programming", "Linear, mixed integer, nonlinear, goal and multiobjective models"],
        ["M-09 to M-12", "Network and routing", "Flows, vehicle routing, multimodal paths, location allocation"],
        ["M-13 to M-17", "Inventory scheduling and queues", "Safety stock, allocation, sequencing, simulation and queueing"],
        ["M-18 to M-21", "Risk probability and decision analysis", "Reliability, stochastic risk, Monte Carlo, decision trees"],
        ["M-22 to M-25", "Robustness decomposition and search", "Robust optimization, decomposition, heuristics, metaheuristics"],
        ["M-26 to M-30", "Learning control and validation", "Reinforcement learning, control, digital twins, causal validation, human factors"],
    ], [1.15, 2.2, 3.5], body_size=8.1)

    page_break(doc)
    add_label(doc, "End to end story")
    add_heading(doc, "Walk Anode Shield from evidence to review", intro="Use this storyline when demonstrating how one project combines apps, agents, OR methods, evidence, and human expertise.")
    add_table(doc, ["Step", "Action", "Narration", "Visible proof"], [
        ["1", "Select Mobility and EV, Apex Mobility, Anode Shield", "The project owns the decision memory and access context.", "P-001 header and project metrics"],
        ["2", "Open Value at risk", "The value is a simulated P90 contribution exposure, not a finance-system fact.", "EV-001-01 receipt"],
        ["3", "Open Data & graph, then Knowledge graph", "The trace connects project evidence to supplier, material, route, plant, product, and decision nodes.", "Selectable nodes and trace steps"],
        ["4", "Open Playground and select the OR specialist", "The visible trace is a product execution log, not private reasoning.", "Agent profile and prompt-bound trace"],
        ["5", "Steer with Preserve service floor and Require human review", "The instruction becomes part of the current trace context.", "Steering chips and trace state"],
        ["6", "Open Network Optimizer", "The app carries the same project and selected decision context into formulation and scenario comparison.", "Project binding strip and command bar"],
        ["7", "Inspect method, data, KPI, and change receipts", "Every app-level count and target says whether it is a catalog item, target contract, or fixture result.", "Session receipt ledger"],
        ["8", "Open Decisions and create expert review draft", "The package binds D0 to D3 choices, L2 L1 L0 variables, methods, and evidence.", "Review draft receipt"],
        ["9", "Open Decision and scenario receipts", "The decision compares alternatives without calling a solver or labeling a result optimal.", "Scenario receipt and claim boundary"],
        ["10", "Open Review", "The release either changes the browser-session decision stage or explains the failed prerequisite.", "Saved or Blocked receipt"],
    ], [0.45, 2.0, 3.1, 1.3], body_size=7.8)
    add_heading(doc, "A second storyline", level=2)
    p = doc.add_paragraph("For ports and cargo, select P-008 Global Berth to Door. Open LogisticsRadar, choose route and transfer layers, focus a port region, pin a cargo or handoff, and inspect quantity and value. Return to Data & graph, use Sources, and request a future QR custody feed or port update under the client and project boundary.")

    page_break(doc)
    add_label(doc, "Deep links and controls")
    add_heading(doc, "Open a precise workspace", intro="Maya records validated project context in the query string so a walkthrough can start at a project surface, Playground session, app run, or mounted app studio.")
    add_code_block(doc, [
        "?view=company&scope=company&sector=mobility-ev&client=apex-mobility&project=anode-shield&projectTab=overview",
        "?view=company&scope=company&sector=mobility-ev&client=apex-mobility&project=anode-shield&projectTab=agents&session=SES-P001-024",
        "?view=company&scope=company&sector=mobility-ev&client=apex-mobility&project=anode-shield&projectTab=apps&run=APP-P001-NO-019",
        "?view=company&scope=company&sector=critical-minerals&client=terrametals&project=lithium-cell-provenance&projectTab=graph",
        "?view=company&scope=company&sector=ports-maritime&client=blueharbor&project=berth-to-door&projectTab=apps&projectApp=logistics",
        "?view=optimizer&scope=company&sector=energy-grid&client=solara-grid&project=copper-rare-earth",
    ])
    add_heading(doc, "Keyboard and shell controls", level=2)
    add_table(doc, ["Control", "Behavior"], [
        ["Command K or Control K", "Open global search"],
        ["Command B or Control B", "Collapse or expand the project rail when focus is outside an input"],
        ["Escape", "Close search, notifications, profile menu, mobile navigation, accountability panel, receipts, evidence, onboarding, or the agent builder"],
        ["Browser Back and Forward", "Restore validated views, scope, decision, project, surface, app studio, Playground session, and app run"],
        ["Tab and Shift Tab", "Move through buttons, inputs, selectors, map entities, and dialog controls; focus remains inside an open modal until it closes"],
        ["Enter in a Playground prompt", "Send the prompt; use Shift Enter to add a new line"],
        ["Enter or Space", "Activate the focused button or map entity"],
        ["Arrow keys on map", "Pan the focused radar viewport"],
        ["Plus and Minus on map", "Zoom the focused radar viewport"],
    ], [2.05, 4.8], body_size=8.4)
    add_heading(doc, "If a control appears not to work", level=2)
    add_bullets(doc, [
        "Look for an inspector, drawer, action receipt, toast, route change, selected state, changed count, or updated chart rather than assuming the control opens a new page.",
        "Read the gate explanation. Disabled or blocked actions require an earlier stage, evidence threshold, feasible candidate, or named approval.",
        "Clear a narrow search or filter when no records are visible.",
        "Reset the graph or map when the current layer combination hides the selected entity.",
        "Refresh only when you want to reset browser-session state. There is no backend persistence in this concept.",
    ])

    page_break(doc)
    add_label(doc, "Production roadmap")
    add_heading(doc, "Concept boundaries and production requirements", intro="The experience demonstrates product architecture and interaction contracts. The following capabilities need production services and verification before operational use.")
    add_table(doc, ["Capability", "Concept status", "Production requirement"], [
        ["Tenant and project security", "UI context only", "Identity provider, server authorization, row and object security, key management, audit"],
        ["Project persistence", "Browser memory", "Project registry, database, object storage, event history, retention and deletion"],
        ["Data ingestion", "File-name and adapter walkthrough", "Upload service, connectors, malware scan, schema and quality services, data contracts"],
        ["Data and knowledge graph", "File-name, connector, query, and deterministic graph fixtures", "Upload and query services, entity resolution, temporal graph, provenance store, access-aware queries, versioning"],
        ["Playground and agents", "Prompt-bound synthetic playback", "Model runtime, evaluated Skills, approved MCP connections, tool sandbox, observability, policy"],
        ["Optimization", "Deterministic response calculator", "Solver adapters, status, incumbent, bounds, gap, infeasibility diagnosis, validation, fallback"],
        ["Learning", "Session receipts only", "Outcome capture, offline evaluation, safety constraints, drift monitoring, controlled promotion and rollback"],
        ["World data", "Local basemap and synthetic records", "Licensed tiles and feeds, attribution, provider compliance, freshness, privacy and resilience"],
        ["Operational write-back", "Not connected", "Human authority, dual control where required, idempotency, transaction receipt, reconciliation and rollback"],
        ["Accessibility and QA", "Automated code checks and static validation", "Browser, screen reader, keyboard, responsive, performance, security and user acceptance testing"],
    ], [1.48, 1.45, 3.95], body_size=7.9)
    add_heading(doc, "Recommended production sequence", level=2)
    add_bullets(doc, [
        "Implement identity, tenant and project authorization, persistence, and audit before connecting client data.",
        "Build the evidence receipt and lineage service as a shared platform contract.",
        "Connect one client project and two apps end to end, with read-only adapters and human review.",
        "Add a solver service with reproducibility and claim-integrity controls before using optimization outputs operationally.",
        "Add geospatial and edge data providers only after licensing, attribution, privacy, and retention decisions are approved.",
        "Introduce outcome learning in shadow mode, then promote agents or policies only through evaluated and reversible releases.",
    ], numbered=True)

    page_break(doc)
    add_label(doc, "Technical references")
    add_heading(doc, "Standards and provider documentation", intro="These references informed the future-state contracts described in the experience. The current concept does not connect the services.")
    references = [
        ("Google Maps Platform 3D Tiles", "https://developers.google.com/maps/documentation/tile/3d-tiles"),
        ("Google Maps Platform policies", "https://developers.google.com/maps/documentation/tile/policies"),
        ("Cesium learning center", "https://cesium.com/learn"),
        ("OGC 3D Tiles standard", "https://www.ogc.org/standards/3dtiles/"),
        ("OpenSky Network terms of use", "https://opensky-network.org/about/terms-of-use"),
        ("OpenSky Network FAQ", "https://opensky-network.org/about/faq"),
        ("ADS-B Exchange data use policy", "https://support.adsbexchange.com/hc/en-us/articles/37364077703693-What-is-ADS-B-Exchange-s-data-use-policy"),
        ("Model Context Protocol specification", "https://modelcontextprotocol.io/specification/2025-11-25/server"),
        ("Claude Code operating model", "https://code.claude.com/docs/en/how-claude-code-works"),
        ("W3C PROV-O", "https://www.w3.org/TR/prov-o/"),
        ("OpenLineage object model", "https://openlineage.io/docs/spec/object-model/"),
        ("OpenTelemetry traces", "https://opentelemetry.io/docs/concepts/signals/traces/"),
        ("NIST AI Risk Management Framework core", "https://airc.nist.gov/airmf-resources/airmf/5-sec-core/"),
        ("Google OR-Tools CP-SAT", "https://developers.google.com/optimization/cp/cp_solver"),
        ("USGS Mineral Commodity Summaries 2026", "https://pubs.usgs.gov/publication/mcs2026"),
    ]
    for index, (label, url) in enumerate(references, 1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.16)
        paragraph.paragraph_format.first_line_indent = Inches(-0.16)
        set_font(paragraph.add_run(f"{index}. "), name="Consolas", size=8.5, bold=True, color=MUTED)
        add_hyperlink(paragraph, label, url)
        set_font(paragraph.add_run(f"  {url}"), name="Consolas", size=7.4, color=MUTED)

    add_heading(doc, "Final demonstration reminder", level=2)
    p = doc.add_paragraph("Use the concept as one coherent workspace: Operations World for global or regional awareness, then the selected Client, Tower, and Project path for delivery. Use the mounted-work bar for apps and accountable people; use Data & graph for sources and knowledge; use Playground for agent-first execution. Keep every decision, variable, method, receipt, and result inside the selected project. When a screen shows a fixture, target contract, missing connector, or blocked gate, preserve that label; it is part of the trust model.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_document()
